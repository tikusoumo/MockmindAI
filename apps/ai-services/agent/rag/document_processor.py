"""Document processor for RAG pipeline.

Handles PDF, TXT, and DOCX document loading, chunking, and embedding.
Following python-patterns skill: async patterns for I/O operations.
"""

from __future__ import annotations

import hashlib
import os
from pathlib import Path
from typing import BinaryIO

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain.schema import Document

from .schemas import DocumentChunk, DocumentMetadata, DocumentType, ProcessedDocument


class DocumentProcessor:
    """Process documents for vector storage.
    
    Handles:
    - PDF, TXT, DOCX file loading
    - Text chunking with overlap
    - Metadata extraction
    """

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""],
        )

    async def process_file(
        self,
        file: BinaryIO,
        filename: str,
        doc_type: DocumentType,
        template_id: str | None,
        uploaded_by: str,
    ) -> ProcessedDocument:
        """Process an uploaded file into chunks.
        
        Args:
            file: File-like object
            filename: Original filename
            doc_type: Type of document
            template_id: Associated template ID (None for user uploads)
            uploaded_by: User ID who uploaded
            
        Returns:
            ProcessedDocument with chunks
        """
        # Generate document ID
        content_bytes = file.read()
        file.seek(0)
        doc_id = self._generate_doc_id(content_bytes, filename)
        
        # Create metadata
        metadata = DocumentMetadata(
            id=doc_id,
            name=filename,
            doc_type=doc_type,
            template_id=template_id,
            uploaded_by=uploaded_by,
            file_size=len(content_bytes),
        )
        
        try:
            # Load and chunk document
            documents = await self._load_document(content_bytes, filename)
            chunks = self._chunk_documents(documents, doc_id)
            
            metadata.chunk_count = len(chunks)
            
            return ProcessedDocument(
                metadata=metadata,
                chunks=chunks,
                status="success",
            )
        except Exception as e:
            return ProcessedDocument(
                metadata=metadata,
                chunks=[],
                status="failed",
                error=str(e),
            )

    async def _load_document(
        self,
        content: bytes,
        filename: str,
    ) -> list[Document]:
        """Load document based on file extension."""
        ext = Path(filename).suffix.lower()
        
        # Save to temp file for loaders that need file path
        temp_path = Path(f"/tmp/{filename}")
        temp_path.write_bytes(content)
        
        try:
            if ext == ".pdf":
                loader = PyPDFLoader(str(temp_path))
                return loader.load()
            elif ext == ".txt" or ext == ".md":
                loader = TextLoader(str(temp_path))
                return loader.load()
            elif ext == ".docx":
                # Use python-docx for DOCX
                from docx import Document as DocxDocument
                doc = DocxDocument(temp_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                return [Document(page_content=text, metadata={"source": filename})]
            else:
                # Fallback: try as text
                text = content.decode("utf-8", errors="ignore")
                return [Document(page_content=text, metadata={"source": filename})]
        finally:
            # Cleanup temp file
            if temp_path.exists():
                temp_path.unlink()

    def _chunk_documents(
        self,
        documents: list[Document],
        doc_id: str,
    ) -> list[DocumentChunk]:
        """Split documents into chunks."""
        chunks = []
        split_docs = self.text_splitter.split_documents(documents)
        
        for i, doc in enumerate(split_docs):
            chunk = DocumentChunk(
                id=f"{doc_id}_chunk_{i}",
                content=doc.page_content,
                metadata={
                    "doc_id": doc_id,
                    "chunk_index": i,
                    "source": doc.metadata.get("source", ""),
                    "page": doc.metadata.get("page", 0),
                },
            )
            chunks.append(chunk)
        
        return chunks

    def _generate_doc_id(self, content: bytes, filename: str) -> str:
        """Generate unique document ID from content hash."""
        hasher = hashlib.sha256()
        hasher.update(content)
        hasher.update(filename.encode())
        return f"doc_{hasher.hexdigest()[:16]}"
